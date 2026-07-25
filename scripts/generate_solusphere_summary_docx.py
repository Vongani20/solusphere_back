"""Generate Solusphere summary and value-add Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = "docs/Solusphere_Summary_and_Value_Add.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 96, 167)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Solusphere", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(75, 0, 130)

    subtitle = doc.add_paragraph("Summary & Value Add")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 151, 167)

    doc.add_paragraph()

    add_heading(doc, "What It Is", level=1)
    doc.add_paragraph(
        "Solusphere is SoluGrowth's integrated workforce and operations platform. The backend "
        "powers a secure web application for employees, managers, and administrators working "
        "in BPO, ITO, and KPO environments. It brings together identity management, collaboration, "
        "AI-assisted support, talent management, and document intelligence in a single API-driven platform."
    )

    add_heading(doc, "Platform Summary", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "What It Does"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    rows = [
        (
            "Identity & Access",
            "Email/password and Outlook 365 signup, JWT sessions, password reset via SMS, "
            "and face recognition login using AWS Rekognition.",
        ),
        (
            "Profile & Security",
            "User profiles, password updates, and mandatory face registration after first login.",
        ),
        (
            "Collaboration",
            "Company events, event chat and comments, and direct messaging between users.",
        ),
        (
            "Helpdesk",
            "Ticket submission and tracking, plus AI-assisted helpdesk chat for instant support.",
        ),
        (
            "AI Assistant (SIA)",
            "General-purpose chatbot with optional web search for operational questions.",
        ),
        (
            "BPO Document Analysis",
            "Upload PDF documents for AI-powered analysis and structured business insights.",
        ),
        (
            "CV Builder",
            "Multi-step CV creation, profile photo upload, SoluGrowth-branded PDF download, "
            "and talent search by skill or qualification.",
        ),
        (
            "Admin Console",
            "Full management of users, roles, events, helpdesk tickets, and employee CVs "
            "(create, view, edit, delete, download).",
        ),
        (
            "File Storage",
            "Secure file uploads to AWS S3 for faces, CV photos, and documents.",
        ),
    ]
    for area, desc in rows:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = desc

    doc.add_paragraph()
    tech = doc.add_paragraph()
    tech.add_run("Technology stack: ").bold = True
    tech.add_run(
        "Go (Gin), MySQL, AWS (S3, Rekognition, SNS, Secrets Manager), OpenAI, "
        "Docker on EC2, CloudFront frontend."
    )

    add_heading(doc, "Value Add", level=1)

    add_heading(doc, "For Employees", level=2)
    add_bullet(doc, " — events, chat, support, and CV building without switching tools.", "One place to work")
    add_bullet(doc, " — face recognition reduces friction and password dependency.", "Faster login")
    add_bullet(doc, " — branded SoluGrowth PDFs with photo, ready for client or internal use.", "Professional CVs")
    add_bullet(doc, " — AI chatbot and helpdesk assistant for common questions.", "Instant help")

    add_heading(doc, "For Managers & HR", level=2)
    add_bullet(doc, " — search CVs by skill or qualification; view and download any employee CV.", "Talent visibility")
    add_bullet(doc, " — admin-created users, role assignment, and face-enforced access.", "Controlled onboarding")
    add_bullet(doc, " — create events and monitor participation and discussion.", "Event engagement")

    add_heading(doc, "For Operations & BPO Teams", level=2)
    add_bullet(doc, " — PDF upload and AI analysis for process review and decision support.", "Document intelligence")
    add_bullet(doc, " — ticket lifecycle from submission to admin resolution.", "Structured helpdesk")
    add_bullet(doc, " — centralized user, ticket, and CV data in one database.", "Audit-ready")

    add_heading(doc, "For the Organisation (SoluGrowth)", level=2)
    add_bullet(doc, " — CVs, logos, and templates aligned with SoluGrowth identity.", "Brand consistency")
    add_bullet(
        doc,
        " — JWT authentication, face verification, private S3 storage, secrets in AWS Secrets Manager.",
        "Security by design",
    )
    add_bullet(doc, " — API-first architecture, Swagger documentation, cloud-ready deployment.", "Scalable foundation")
    add_bullet(
        doc,
        " — OpenAI integrated into support, chat, and BPO analysis workflows.",
        "AI differentiation",
    )

    add_heading(doc, "One-Line Pitch", level=1)
    pitch = doc.add_paragraph()
    pitch.add_run(
        '"Solusphere unifies secure workforce access, collaboration, AI-powered support, '
        'and branded talent management into a single platform built for SoluGrowth\'s BPO operations."'
    ).italic = True

    add_heading(doc, "Suggested Use Cases", level=1)
    use_cases = [
        "New hire onboarding — register, complete face enrolment, build CV, download branded PDF.",
        "Client staffing — admin searches talent by skill, reviews CVs, downloads PDFs for proposals.",
        "Internal support — employee raises helpdesk ticket or uses AI chat for quick answers.",
        "Process improvement — team uploads BPO PDFs for AI analysis and insight extraction.",
        "Company communications — admin creates events; staff join and discuss in real time.",
    ]
    for i, case in enumerate(use_cases, 1):
        doc.add_paragraph(f"{i}. {case}", style="List Number")

    doc.add_paragraph()
    footer = doc.add_paragraph("Prepared for SoluGrowth | Solusphere Backend Platform")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
