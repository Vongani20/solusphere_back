"""Generate Solusphere system usage guide Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = "docs/Solusphere_System_Usage_Guide.docx"

TEAL = RGBColor(0, 151, 167)
PURPLE = RGBColor(75, 0, 130)
GRAY = RGBColor(100, 100, 100)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL if level > 1 else PURPLE
    return h


def steps(doc, items):
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def note(doc, text):
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run(text)


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Solusphere System Usage Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PURPLE

    sub = doc.add_paragraph("How to use the platform — employees, managers, and administrators")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = TEAL
    doc.add_paragraph()

    # 1. Overview
    heading(doc, "1. Overview", 1)
    doc.add_paragraph(
        "Solusphere is SoluGrowth's workforce platform, accessed through a web browser. "
        "Users sign in securely, complete face registration once, and then use collaboration, "
        "support, AI, and CV tools from a single application."
    )
    bullets(doc, [
        "Web application (frontend): hosted on CloudFront / S3",
        "API (backend): secure REST API with JWT authentication",
        "API documentation: /swagger/index.html on the backend server",
    ])

    # 2. User roles
    heading(doc, "2. User Roles", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "What they can do"
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    roles = [
        ("Employee (standard user)", "Profile, face login, events, chat, helpdesk, AI assistant, CV Builder, BPO analysis"),
        ("Administrator", "Everything an employee can do, plus user management, events, helpdesk tickets, and CV oversight"),
    ]
    for role, desc in roles:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = desc
    doc.add_paragraph()

    # 3. Getting started
    heading(doc, "3. Getting Started", 1)

    heading(doc, "3.1 Create an account", 2)
    steps(doc, [
        "Open the Solusphere web application in your browser.",
        "Choose Register (local account) or Create Outlook 365 Account.",
        "Enter username, email, phone number, and password.",
        "Submit the form. You will be prompted to log in next.",
    ])
    note(doc, "Phone number is used for password reset SMS codes.")

    heading(doc, "3.2 First login and face registration (required)", 2)
    doc.add_paragraph(
        "After your first password login, you must register your face before using most features. "
        "This is a one-time security step."
    )
    steps(doc, [
        "Log in with your email and password.",
        "Go to Profile or follow the on-screen prompt to register your face.",
        "Allow camera access when prompted.",
        "Capture a clear front-facing photo in good lighting.",
        "Submit face registration. Once confirmed, all app features unlock.",
    ])
    note(doc, "Until face registration is complete, protected features return a 'Face registration required' message.")

    heading(doc, "3.3 Everyday login options", 2)
    bullets(doc, [
        "Password login — email and password on the Sign In page.",
        "Face login — use the Face Login option and submit a live face image.",
    ])

    heading(doc, "3.4 Forgot password", 2)
    steps(doc, [
        "Click Forgot? on the Sign In page.",
        "Enter your registered email address.",
        "Receive a reset code by SMS and email.",
        "Enter the code and your new password on the reset screen.",
        "Log in with the new password.",
    ])

    # 4. Employee features
    heading(doc, "4. Employee Features", 1)

    heading(doc, "4.1 Profile", 2)
    bullets(doc, [
        "View your account details from the Profile page.",
        "Update your password (requires current password).",
        "Update or delete your registered face if needed.",
    ])

    heading(doc, "4.2 CV Builder", 2)
    doc.add_paragraph(
        "The CV Builder is a four-step wizard. Your progress is saved as you move between steps."
    )
    steps(doc, [
        "Step 1 — Personal Information: name, gender, nationality, date of birth, profile summary, value proposition.",
        "Step 2 — Skills & Qualifications: professional skills (with detail points), qualifications, computer skills, memberships, languages.",
        "Step 3 — Experience: company, position, period, and scope of work for each role.",
        "Step 4 — Review & Photo: upload a profile photo (JPG or PNG, max 5 MB), review all sections, then Save.",
        "Download CV: click Download to generate a two-page SoluGrowth-branded PDF with your photo and logo.",
    ])
    note(doc, "You can return later to edit your CV. Use Delete CV only if you want to remove your entire CV record.")

    heading(doc, "4.3 Talent search", 2)
    doc.add_paragraph(
        "Authenticated users can search the talent directory by skill or qualification keyword "
        "(e.g. SAP, BCom). Results show matching employee CV summaries."
    )

    heading(doc, "4.4 Events", 2)
    steps(doc, [
        "Open Events to see company events created by administrators.",
        "Select an event and click Join to participate.",
        "View and post messages or comments in the event chat.",
        "Engage with colleagues on event images and discussions.",
    ])

    heading(doc, "4.5 Direct messaging", 2)
    steps(doc, [
        "Open Chats to see your conversations.",
        "Select a colleague from the user list.",
        "Read message history and send new direct messages.",
    ])

    heading(doc, "4.6 Helpdesk", 2)
    bullets(doc, [
        "Submit a ticket — describe your issue; it is logged for the support team.",
        "Helpdesk AI chat — get instant AI-assisted answers for common support questions.",
    ])

    heading(doc, "4.7 AI Assistant (SIA)", 2)
    doc.add_paragraph(
        "Use SIA Chat to ask operational questions, search the web, review websites, "
        "and request analytics summaries. SIA can also generate downloadable research reports."
    )
    bullets(doc, [
        "Chat: ask questions in natural language; web search is enabled by default.",
        "Reports: request a research report and choose Word, Excel, or PowerPoint output.",
        "Example prompts: market research, competitor analysis, website review, industry trends.",
    ])
    steps(doc, [
        "Open SIA Chat from the sidebar.",
        "For chat answers, type your question and send.",
        "For a report, request export in Word (.docx), Excel (.xlsx), or PowerPoint (.pptx) format.",
        "Download the generated file when SIA finishes researching your topic.",
    ])

    heading(doc, "4.8 BPO document analysis", 2)
    steps(doc, [
        "Open BPO Analysis in the application.",
        "Upload a PDF document.",
        "Wait for processing to complete.",
        "Review the AI-generated analysis results.",
        "Access past analyses from your history; delete when no longer needed.",
    ])

    # 5. Administrator features
    heading(doc, "5. Administrator Features", 1)
    doc.add_paragraph(
        "Administrators have an extended Admin section. The first user to run Admin Bootstrap "
        "on a new deployment is promoted to admin."
    )

    heading(doc, "5.1 User management", 2)
    bullets(doc, [
        "List all users and view individual profiles.",
        "Create new users on behalf of the organisation.",
        "Update user details or change roles (employee / admin).",
        "Delete users when they leave the organisation.",
    ])

    heading(doc, "5.2 Event management", 2)
    bullets(doc, [
        "Create events with title, description, and optional image.",
        "Edit or delete existing events.",
        "Monitor and participate in event chat as an admin.",
    ])

    heading(doc, "5.3 Helpdesk management", 2)
    steps(doc, [
        "View all submitted helpdesk tickets.",
        "Open a ticket to read full details.",
        "Update status, subject, or description.",
        "Close or delete resolved tickets.",
    ])

    heading(doc, "5.4 CV management (HR / talent)", 2)
    steps(doc, [
        "Open Admin → CVs to list all employee CV profiles.",
        "Filter by skill or qualification (e.g. Python, MBA).",
        "Select a user to view their full CV data.",
        "Download any employee's branded SoluGrowth PDF for client proposals or internal review.",
    ])

    # 6. Typical workflows
    heading(doc, "6. Typical Day-to-Day Workflows", 1)

    workflows = [
        (
            "New employee first day",
            "Register → log in → register face → complete CV Builder → download branded PDF.",
        ),
        (
            "Daily login",
            "Open app → Face Login (or password) → access dashboard features.",
        ),
        (
            "Need IT or HR help",
            "Submit helpdesk ticket OR use Helpdesk AI chat for quick guidance.",
        ),
        (
            "Company announcement",
            "Admin creates event → employees join → discussion in event chat.",
        ),
        (
            "Staffing a client project",
            "Admin searches CVs by skill → reviews candidates → downloads PDFs for proposal.",
        ),
        (
            "Process review",
            "Upload BPO PDF → review AI analysis → share insights with team.",
        ),
    ]
    for name, flow in workflows:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(flow)

    # 7. Requirements
    heading(doc, "7. System Requirements", 1)
    bullets(doc, [
        "Modern web browser (Chrome, Edge, Firefox, or Safari).",
        "Stable internet connection.",
        "Webcam for face registration and face login.",
        "Valid email and mobile number on your account.",
        "For CV photo upload: JPG or PNG image, maximum 5 MB.",
    ])

    # 8. Security
    heading(doc, "8. Security & Access Rules", 1)
    bullets(doc, [
        "All protected features require a valid login session (JWT token).",
        "Face registration is mandatory before using most features after first login.",
        "Only administrators can manage users, events, tickets, and all employee CVs.",
        "Users can only view, edit, and delete their own CV.",
        "Files (faces, photos, documents) are stored securely in AWS S3.",
        "Passwords are never stored in plain text.",
    ])

    # 9. Support
    heading(doc, "9. Getting Help", 1)
    bullets(doc, [
        "Use in-app Helpdesk or Helpdesk AI chat for platform issues.",
        "Contact your SoluGrowth administrator for account or access problems.",
        "Developers can refer to Swagger API docs at /swagger/index.html for technical integration.",
    ])

    doc.add_paragraph()
    footer = doc.add_paragraph("SoluGrowth | Solusphere Platform — System Usage Guide")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
